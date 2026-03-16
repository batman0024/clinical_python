
from typing import List, Dict
from docx import Document


def parse_shells_docx(path) -> List[Dict]:
    """
    Extract tables/text blocks that look like shells (Table titles, endpoints, footnotes).
    Minimal heuristic: read tables and paragraphs; mark 'Table'/'Figure' headers.
    """
    doc = Document(str(path))
    entries = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt and ("Table" in txt or "Figure" in txt or "Listing" in txt):
            entries.append({"type": "shell_header", "text": txt})
    for t in doc.tables:
        rows = []
        for row in t.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        entries.append({"type": "shell_table", "rows": rows})
    return entries
